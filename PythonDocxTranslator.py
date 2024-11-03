# dont forget to install dependencies python-docx  'pip install requests python-docx'
import requests
from docx import Document
import os

subscription_key = "YOUR_SUBSCRIPTION_KEY"
endpoint = 'https://api.cognitive.microsofttranslator.com'
location = 'YOUR_RESOUCE_LOCATION'
#Here i set pt-br but you can change based os you destination language
language_destination = 'pt-br'

def translator_text(text, target_language):
  path = '/translate'
  constructed_url = endpoint + path
  headers = {
      'Ocp-Apim-Subscription-Key': subscription_key,
      'Ocp-Apim-Subscription-Region': location,
      'Content-Type': 'application/json',
      'X-clientTraceId': str(os.urandom(16))      
  }

  body = [{
      'text': text
  }]
  params = {
      'api-version': '3.0',
      'from': 'en',
      'to': [target_language]
  }
  request = requests.post(constructed_url, params=params, headers=headers, json=body)
  response = request.json()
  return response[0]["translations"][0]["text"]

def translate_document(path):
  document = Document(path)
  full_text = []
  for paragraph in document.paragraphs:
    translated_text = translator_text(paragraph.text, language_destination)
    full_text.append(translated_text)

  translated_doc = Document()
  for line in full_text:
    print(line)
    translated_doc.add_paragraph(line)
  path_translated = path.replace(".docx", f"_{language_destination}.docx")
  translated_doc.save(path_translated)
  return path_translated

input_file = "/content/thegirl.docx"
translate_document(input_file)